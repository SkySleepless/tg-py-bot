"""
Content extractor for extracting text from various file formats.
Supports images (OCR), PDFs, DOCX, and plain text files.
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Optional, Dict, Any
import io

logger = logging.getLogger(__name__)


class ContentExtractor:
    """Extracts text content from various file formats."""
    
    def __init__(self):
        """Initialize the content extractor."""
        self.ocr_available = self._check_ocr_availability()
    
    def _check_ocr_availability(self) -> bool:
        """Check if OCR dependencies are available."""
        try:
            import pytesseract
            from PIL import Image
            return True
        except ImportError:
            logger.warning("OCR dependencies not available. Install: pip install pytesseract pillow")
            return False
    
    def extract_text(self, file_path: str, file_category: str) -> Dict[str, Any]:
        """
        Extract text from a file based on its category.
        
        Args:
            file_path: Path to the file
            file_category: File category ('image' or 'document')
            
        Returns:
            Dictionary with extraction results
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "content": "",
                "metadata": {}
            }
        
        try:
            if file_category == "image":
                return self._extract_from_image(file_path)
            elif file_category == "document":
                return self._extract_from_document(file_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file category: {file_category}",
                    "content": "",
                    "metadata": {}
                }
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def _extract_from_image(self, image_path: str) -> Dict[str, Any]:
        """
        Extract text from an image using OCR.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Dictionary with extraction results
        """
        if not self.ocr_available:
            return {
                "success": False,
                "error": "OCR not available. Install pytesseract and pillow.",
                "content": "",
                "metadata": {}
            }
        
        try:
            import pytesseract
            from PIL import Image
            
            # Open and process image
            image = Image.open(image_path)
            
            # Convert to grayscale for better OCR
            if image.mode != 'L':
                image = image.convert('L')
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image)
            
            # Clean up text
            text = text.strip()
            
            metadata = {
                "format": image.format,
                "size": image.size,
                "mode": image.mode,
                "ocr_engine": "Tesseract"
            }
            
            return {
                "success": True,
                "content": text,
                "metadata": metadata,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"OCR failed for {image_path}: {e}")
            return {
                "success": False,
                "error": f"OCR failed: {str(e)}",
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def _extract_from_document(self, document_path: str) -> Dict[str, Any]:
        """
        Extract text from a document (PDF, DOCX, TXT, etc.).
        
        Args:
            document_path: Path to the document file
            
        Returns:
            Dictionary with extraction results
        """
        file_path = Path(document_path)
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._extract_from_pdf(document_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(document_path)
            elif ext in ['.txt', '.md', '.rtf']:
                return self._extract_from_text(document_path)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported document format: {ext}",
                    "content": "",
                    "metadata": {}
                }
        except Exception as e:
            logger.error(f"Document extraction failed for {document_path}: {e}")
            return {
                "success": False,
                "error": f"Document extraction failed: {str(e)}",
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def _extract_from_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """
        Extract text from a PDF file.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary with extraction results
        """
        try:
            # Try pdfplumber first (better for complex PDFs)
            try:
                import pdfplumber
                text_parts = []
                metadata = {"library": "pdfplumber"}
                
                with pdfplumber.open(pdf_path) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                
                text = "\n\n".join(text_parts)
                if text.strip():
                    return {
                        "success": True,
                        "content": text,
                        "metadata": metadata,
                        "error": None
                    }
            except ImportError:
                logger.debug("pdfplumber not available, trying PyPDF2")
            
            # Fallback to PyPDF2
            try:
                import PyPDF2
                text_parts = []
                metadata = {"library": "PyPDF2"}
                
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                
                text = "\n\n".join(text_parts)
                return {
                    "success": True,
                    "content": text,
                    "metadata": metadata,
                    "error": None
                }
            except ImportError:
                logger.debug("PyPDF2 not available")
            
            # No PDF library available
            return {
                "success": False,
                "error": "No PDF library available. Install pdfplumber or PyPDF2.",
                "content": "",
                "metadata": {}
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed for {pdf_path}: {e}")
            return {
                "success": False,
                "error": f"PDF extraction failed: {str(e)}",
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def _extract_from_docx(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary with extraction results
        """
        try:
            import docx
            
            doc = docx.Document(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            metadata = {
                "library": "python-docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return {
                "success": True,
                "content": text,
                "metadata": metadata,
                "error": None
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "python-docx not available. Install: pip install python-docx",
                "content": "",
                "metadata": {}
            }
        except Exception as e:
            logger.error(f"DOCX extraction failed for {docx_path}: {e}")
            return {
                "success": False,
                "error": f"DOCX extraction failed: {str(e)}",
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def _extract_from_text(self, text_path: str) -> Dict[str, Any]:
        """
        Extract text from a plain text file.
        
        Args:
            text_path: Path to the text file
            
        Returns:
            Dictionary with extraction results
        """
        try:
            with open(text_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            metadata = {
                "encoding": "utf-8",
                "lines": len(text.splitlines())
            }
            
            return {
                "success": True,
                "content": text,
                "metadata": metadata,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Text extraction failed for {text_path}: {e}")
            return {
                "success": False,
                "error": f"Text extraction failed: {str(e)}",
                "content": "",
                "metadata": {"exception": type(e).__name__}
            }
    
    def cleanup_file(self, file_path: str) -> bool:
        """
        Clean up (delete) a file after processing.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file was deleted successfully
        """
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.debug(f"Cleaned up file: {file_path}")
                return True
            return False
        except Exception as e:
            logger.warning(f"Failed to clean up file {file_path}: {e}")
            return False


# Singleton instance
_content_extractor: Optional[ContentExtractor] = None

def get_content_extractor() -> ContentExtractor:
    """
    Get or create the content extractor singleton.
    
    Returns:
        ContentExtractor instance
    """
    global _content_extractor
    if _content_extractor is None:
        _content_extractor = ContentExtractor()
    return _content_extractor